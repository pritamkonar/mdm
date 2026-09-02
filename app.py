import streamlit as st
import pandas as pd
import io
import json
import urllib.request
import zipfile
import os
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.worksheet.page import PageMargins
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Streamlit Page Config & Custom Styling
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="PM POSHAN / MDM Monthly Report & UC Generator",
    page_icon="🍲",
    layout="wide"
)

st.markdown("""
<style>
    .reportview-container { background: #fdfdfd; }
    .metric-card {
        background-color: #f8f9fa;
        border-radius: 8px;
        padding: 12px 16px;
        border-left: 4px solid #1f77b4;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Bengali Phonetic Transliteration (Offline Dictionary + Google API Cache)
# ---------------------------------------------------------------------------
COMMON_BENGALI_DICT = {
    "bhat": "ভাত", "vat": "ভাত", "rice": "ভাত",
    "dal": "ডাল", "daal": "ডাল",
    "alu": "আলু", "aloo": "আলু", "potato": "আলু",
    "dim": "ডিম", "deem": "ডিম", "egg": "ডিম",
    "mach": "মাছ", "maach": "মাছ", "fish": "মাছ",
    "mangsho": "মাংস", "mangso": "মাংস", "meat": "মাংস", "chicken": "চিকেন",
    "potol": "পটল", "patol": "পটল",
    "begun": "বেগুন", "eggplant": "বেগুন",
    "borboti": "বরবটি", "barbati": "বরবটি",
    "soyabin": "সয়াবিন", "soyabean": "সয়াবিন", "sobin": "সয়াবিন",
    "posto": "পোস্ত",
    "vendi": "ভেন্ডি", "bhendi": "ভেন্ডি", "dharosh": "ঢ্যাঁড়শ",
    "kumro": "কুমড়ো", "kumra": "কুমড়ো", "kumro,": "কুমড়ো,",
    "bandhakopi": "বাঁধাকপি", "badhakopi": "বাঁধাকপি", "cabbage": "বাঁধাকপি",
    "fulkopi": "ফুলকপি", "cauliflower": "ফুলকপি",
    "kochu": "কচু", "kachu": "কচু",
    "pui": "পুঁই", "pui shak": "পুঁই শাক", "puishak": "পুঁইশাক",
    "sobji": "সবজি", "sabji": "সবজি",
    "khichuri": "খিচুড়ি", "khichudi": "খিচুড়ি"
}

@st.cache_data(show_spinner=False)
def transliterate_token(token: str) -> str:
    clean_lower = token.lower().strip()
    if clean_lower in COMMON_BENGALI_DICT:
        return COMMON_BENGALI_DICT[clean_lower]
    
    # If already Bengali characters, keep intact
    if any('\u0980' <= c <= '\u09ff' for c in token):
        return token
    
    # Attempt Google Input Tools API
    try:
        url = f"https://inputtools.google.com/request?text={urllib.request.quote(token)}&itc=bn-t-i0-und&num=1"
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=1.5) as response:
            res_data = json.loads(response.read().decode('utf-8'))
            if res_data[0] == "SUCCESS":
                return res_data[1][0][1][0]
    except Exception:
        pass
    return token

def google_phonetic_convert(text):
    if pd.isna(text): return ""
    text_str = str(text).strip()
    if not text_str: return ""
    
    tokens = [t.strip() for t in text_str.replace(",", ", ").split()]
    converted_tokens = []
    
    for token in tokens:
        if not token: continue
        is_comma = token.endswith(",")
        clean_token = token.rstrip(",")
        
        converted = transliterate_token(clean_token)
        if is_comma:
            converted_tokens.append(converted + ",")
        else:
            converted_tokens.append(converted)
            
    return " ".join(converted_tokens)

def ensure_compulsory_menu(text, has_students=False):
    """
    Ensures 'ভাত, ডাল, আলু' (Rice, Dal, Potato) is compulsory on MDM serving days:
    - User does not need to type 'ভাত, ডাল, আলু' every day.
    - If user types just extra vegetables/items (e.g. 'potol', 'dim', 'begun'), it automatically prepends 'ভাত, ডাল, আলু, '.
    - If user enters students but leaves item blank, defaults to 'ভাত, ডাল, আলু'.
    - Prevents duplicates if user already typed 'ভাত, ডাল, আলু'.
    """
    if text is None or str(text).strip() in ['', 'None', 'nan']:
        return 'ভাত, ডাল, আলু' if has_students else ''
    
    converted = google_phonetic_convert(text).strip()
    if not converted:
        return 'ভাত, ডাল, আলু' if has_students else ''
    
    converted = converted.lstrip(',').strip()
    
    # Check if it already starts with the compulsory base
    base_variants = ['ভাত, ডাল, আলু,', 'ভাত, ডাল, আলু ,', 'ভাত, ডাল, আলু']
    for b in base_variants:
        if converted.startswith(b):
            return converted
            
    # If neither 'ভাত' nor 'ডাল' are present, prepend 'ভাত, ডাল, আলু, '
    if 'ভাত' not in converted and 'ডাল' not in converted:
        if converted.startswith('আলু'):
            extra = converted[len('আলু'):].lstrip(',').strip()
            if extra:
                return f'ভাত, ডাল, আলু, {extra}'
            else:
                return 'ভাত, ডাল, আলু'
        return f'ভাত, ডাল, আলু, {converted}'
    
    # If user typed 'bhat, dal, potol' or similar without alu
    if 'ভাত' in converted and 'ডাল' in converted and 'আলু' not in converted:
        idx = converted.find('ডাল')
        before = converted[:idx+len('ডাল')]
        after = converted[idx+len('ডাল'):].lstrip(',').strip()
        if after:
            return f'{before}, আলু, {after}'
        else:
            return f'{before}, আলু'
            
    return converted

# ---------------------------------------------------------------------------
# Custom Roundoff Rule: No decimals except .5
# - If decimal is .5 (e.g. 20.5, 7606.5), keep .5
# - If decimal < 0.5 (e.g. 20.4), round down (20)
# - If decimal > 0.5 (e.g. 20.6), round up (21)
# ---------------------------------------------------------------------------
def round_special_num(val):
    if val is None or str(val).strip() in ['', 'None']:
        return 0
    try:
        f = float(str(val).replace('Rs', '').replace('Kg', '').replace('kg', '').replace('₹', '').replace(',', '').strip())
        int_part = int(f)
        dec_part = round(f - int_part, 4)
        if abs(dec_part - 0.5) < 1e-4:
            return float(f"{int_part}.5")
        elif dec_part < 0.5:
            return int_part
        else:
            return int_part + 1
    except (ValueError, TypeError):
        return val

def format_val_with_unit(val, unit="", round_to_int=False, is_nil_allowed=False):
    """
    Formats numeric values without unit names (plain digits):
    - No decimal number except .5
    - If empty or None, returns 'NIL' if is_nil_allowed else ''
    - If 'NIL', returns 'NIL'
    - If 0 and is_nil_allowed, returns 'NIL'
    - If decimal is .5 (e.g. 20.5, 7606.5), keeps .5
    - If decimal < 0.5 (e.g. 20.4), rounds down to integer (20)
    - If decimal > 0.5 (e.g. 20.6), rounds up to integer (21)
    """
    if val is None or str(val).strip() in ['', 'None']:
        return 'NIL' if is_nil_allowed else ''
    s = str(val).strip()
    if s.upper() == 'NIL':
        return 'NIL'
    try:
        f = float(s.replace('Rs', '').replace('Kg', '').replace('kg', '').replace('₹', '').replace(',', '').strip())
        if f == 0 and is_nil_allowed:
            return 'NIL'
        int_part = int(f)
        dec_part = round(f - int_part, 4)
        if abs(dec_part - 0.5) < 1e-4:
            return f"{int_part}.5"
        elif dec_part < 0.5:
            return str(int_part)
        else:
            return str(int_part + 1)
    except (ValueError, TypeError):
        return s

def to_excel_num(val, is_nil_allowed=False):
    """
    Converts values to proper Excel numeric types (int / float) or 'NIL':
    - Never returns numeric strings, preventing Excel 'Number Stored as Text' green triangles.
    - If empty, None, or 0 and is_nil_allowed: returns 'NIL'
    - If decimal is .5 (e.g. 20.5, 7606.5), returns float (20.5, 7606.5)
    - If decimal < 0.5 (e.g. 20.4), returns int (20)
    - If decimal > 0.5 (e.g. 20.6), returns int (21)
    """
    if val is None or str(val).strip() in ['', 'None']:
        return 'NIL' if is_nil_allowed else 0
    s = str(val).strip()
    if s.upper() == 'NIL':
        return 'NIL'
    try:
        f = float(s.replace('Rs', '').replace('Kg', '').replace('kg', '').replace('₹', '').replace(',', '').strip())
        if f == 0 and is_nil_allowed:
            return 'NIL'
        int_part = int(f)
        dec_part = round(f - int_part, 4)
        if abs(dec_part - 0.5) < 1e-4:
            return float(f"{int_part}.5")
        elif dec_part < 0.5:
            return int_part
        else:
            return int_part + 1
    except (ValueError, TypeError):
        return val

# ---------------------------------------------------------------------------
# Sample Reference Data (August 2026 - Shyambazar D.N. High School)
# ---------------------------------------------------------------------------
REFERENCE_SAMPLE_ROWS = [
    {"Date": "01.08.26", "No of student availing MDM Class - V": 25, "No of student availing MDM Class - VI to VIII": 153, "Items served": "ভাত, ডাল, আলু বেগুন বরবটি"},
    {"Date": "03.08.26", "No of student availing MDM Class - V": 24, "No of student availing MDM Class - VI to VIII": 155, "Items served": "ভাত, ডাল, আলু ,পটল"},
    {"Date": "04.08.26", "No of student availing MDM Class - V": 24, "No of student availing MDM Class - VI to VIII": 153, "Items served": "ভাত, ডাল, আলু ,কচু পুঁই"},
    {"Date": "05.08.26", "No of student availing MDM Class - V": 24, "No of student availing MDM Class - VI to VIII": 151, "Items served": "ভাত, ডাল, আলু ,পোস্ত"},
    {"Date": "06.08.26", "No of student availing MDM Class - V": 23, "No of student availing MDM Class - VI to VIII": 152, "Items served": "ভাত, ডাল, আলু ,ভেন্ডি"},
    {"Date": "07.08.26", "No of student availing MDM Class - V": 27, "No of student availing MDM Class - VI to VIII": 157, "Items served": "ভাত, ডাল, আলু ,কুমড়ো, সয়াবিন"},
    {"Date": "10.08.26", "No of student availing MDM Class - V": 21, "No of student availing MDM Class - VI to VIII": 145, "Items served": "ভাত, ডাল,আলু , পটল"},
    {"Date": "11.08.26", "No of student availing MDM Class - V": 23, "No of student availing MDM Class - VI to VIII": 135, "Items served": "ভাত, ডাল,আলু  বেগুন, বরবটি"},
    {"Date": "12.08.26", "No of student availing MDM Class - V": 23, "No of student availing MDM Class - VI to VIII": 148, "Items served": "ভাত, ডাল,আলু  ,ডিম, ভেন্ডি"},
    {"Date": "13.08.26", "No of student availing MDM Class - V": 22, "No of student availing MDM Class - VI to VIII": 147, "Items served": "ভাত, ডাল, আলু ,পোস্ত"},
    {"Date": "14.08.26", "No of student availing MDM Class - V": 0, "No of student availing MDM Class - VI to VIII": 121, "Items served": "ভাত, ডাল, আলু ,কুমড়ো, সয়াবিন"},
    {"Date": "17.08.26", "No of student availing MDM Class - V": 26, "No of student availing MDM Class - VI to VIII": 175, "Items served": "ভাত, ডাল, আলু ,কচু, পুই"},
    {"Date": "18.08.26", "No of student availing MDM Class - V": 21, "No of student availing MDM Class - VI to VIII": 122, "Items served": "ভাত, ডাল, আলু বরবটি, বেগুন"},
    {"Date": "19.08.26", "No of student availing MDM Class - V": 26, "No of student availing MDM Class - VI to VIII": 158, "Items served": "ভাত, ডাল, আলু ,ভেন্ডি"},
    {"Date": "20.08.26", "No of student availing MDM Class - V": 27, "No of student availing MDM Class - VI to VIII": 169, "Items served": "ভাত, ডাল, আলু ,পটল, ডিম"},
    {"Date": "21.08.26", "No of student availing MDM Class - V": 28, "No of student availing MDM Class - VI to VIII": 162, "Items served": "ভাত, ডাল, আলু ,পোস্ত"},
    {"Date": "22.08.26", "No of student availing MDM Class - V": 25, "No of student availing MDM Class - VI to VIII": 153, "Items served": "ভাত, ডাল, আলু ,কুমড়ো, সয়াবিন"},
    {"Date": "24.08.26", "No of student availing MDM Class - V": 25, "No of student availing MDM Class - VI to VIII": 162, "Items served": "ভাত, ডাল, আলু ,ভেন্ডি"},
    {"Date": "25.08.26", "No of student availing MDM Class - V": 26, "No of student availing MDM Class - VI to VIII": 168, "Items served": "ভাত, ডাল, আলু ,বাঁধাকপি"},
    {"Date": "27.08.26", "No of student availing MDM Class - V": 24, "No of student availing MDM Class - VI to VIII": 156, "Items served": "ভাত, ডাল, আলু ,পোস্ত"},
    {"Date": "29.08.26", "No of student availing MDM Class - V": 26, "No of student availing MDM Class - VI to VIII": 156, "Items served": "ভাত, ডাল, আলু ,পটল"},
    {"Date": "31.08.26", "No of student availing MDM Class - V": 28, "No of student availing MDM Class - VI to VIII": 158, "Items served": "ভাত, ডাল,আলু  কচু, পুই"}
]

COLUMNS = [
    "Date",
    "No of student availing MDM Class - V",
    "No of student availing MDM Class - VI to VIII",
    "Items served",
    "Cooking cost for Class - V to VIII",
    "Food Grains for Class - V (Kg.)",
    "Food grains for Class - VI to VIII (Kg.)"
]

DEFAULT_COOKS = [
    {"S.No.": 1, "Cook Name": "DILIP BAIRAGYA", "Gender (M/F)": "M", "Category (SC/ST/OBC/GEN)": "GEN", "Payment Mode (Cash/Bank)": "BANK", "Amount Received (In Rs.)": "NIL"},
    {"S.No.": 2, "Cook Name": "APARNA GHOSH", "Gender (M/F)": "F", "Category (SC/ST/OBC/GEN)": "GEN", "Payment Mode (Cash/Bank)": "BANK", "Amount Received (In Rs.)": "NIL"},
    {"S.No.": 3, "Cook Name": "DILIP KUMAR BAIRAGYA", "Gender (M/F)": "M", "Category (SC/ST/OBC/GEN)": "GEN", "Payment Mode (Cash/Bank)": "BANK", "Amount Received (In Rs.)": "NIL"},
    {"S.No.": 4, "Cook Name": "RASED MOLLA", "Gender (M/F)": "M", "Category (SC/ST/OBC/GEN)": "MINORITY", "Payment Mode (Cash/Bank)": "BANK", "Amount Received (In Rs.)": "NIL"},
    {"S.No.": 5, "Cook Name": "BASANA LOHAR", "Gender (M/F)": "F", "Category (SC/ST/OBC/GEN)": "SC", "Payment Mode (Cash/Bank)": "BANK", "Amount Received (In Rs.)": "NIL"}
]

if "df" not in st.session_state:
    st.session_state.df = pd.DataFrame(columns=COLUMNS)

if "cooks_df" not in st.session_state:
    st.session_state.cooks_df = pd.DataFrame(DEFAULT_COOKS)

# ---------------------------------------------------------------------------
# Sidebar: School Profile & Rates Configuration
# ---------------------------------------------------------------------------
with st.sidebar:
    st.header("⚙️ School Profile & Settings")
    month_year = st.text_input("Month & Year", value="August 2026")
    school_name = st.text_input("School Name", value="SHYAMBAZAR D N HIGH SCHOOL")
    udise_code = st.text_input("UDISE Code", value="19251804002")
    
    col_e1, col_e2 = st.columns(2)
    with col_e1:
        enrol_v = st.number_input("Enrolment Class V", value=60, step=1)
    with col_e2:
        enrol_vi = st.number_input("Enrolment Class VI-VIII", value=322, step=1)
        
    st.subheader("📍 Location Details")
    state_name = st.text_input("State / UT", value="WEST BENGAL")
    district_name = st.text_input("District", value="PURBA BARDHAMAN")
    block_name = st.text_input("Block / NP", value="MONGALKOTE")
    circle_name = st.text_input("Circle", value="Monglkote-II")
    gram_sansad = st.text_input("Gram Sansad / Ward", value="Bhallagram")
    village_name = st.text_input("Village", value="SHYAMBAZAR")
    
    st.subheader("🏦 SHG & Bank Details")
    shg_bank_info = st.text_area(
        "SHG / Bank String (Printed on row 41)",
        value="Name of SHG (Cook-cum-Helper) MC - APPROVED_Name of Bank with Branch CANARA BANK, Mathrun Branch._ A/c no_3641101002012_worked during month.......",
        height=80
    )
    
    st.markdown("---")
    st.subheader("💰 Per-Student Rates")
    cost_rate_v = st.number_input("Class V Cost (Rs/stu)", value=10.00, step=0.01, format="%.2f")
    cost_rate_vi = st.number_input("Class VI-VIII Cost (Rs/stu)", value=10.17, step=0.01, format="%.2f")
    grain_rate_v = st.number_input("Class V Grains (Kg/stu)", value=0.100, step=0.005, format="%.3f")
    grain_rate_vi = st.number_input("Class VI-VIII Grains (Kg/stu)", value=0.150, step=0.005, format="%.3f")

# ---------------------------------------------------------------------------
# Main Page Header & Quick Controls
# ---------------------------------------------------------------------------
st.title("🍲 PM POSHAN / MDM Monthly Report & UC Generator")
st.write(f"Generating reports for **{school_name}** ({month_year}) | UDISE: `{udise_code}`")

btn_c1, btn_c2, btn_c3 = st.columns([2, 2, 6])
with btn_c1:
    if st.button("📋 Load August 2026 Sample", type="secondary", help="Populates table and balances from the reference sample"):
        sample_df = pd.DataFrame(REFERENCE_SAMPLE_ROWS)
        sample_df["Cooking cost for Class - V to VIII"] = [
            round_special_num((r["No of student availing MDM Class - V"] * cost_rate_v) + (r["No of student availing MDM Class - VI to VIII"] * cost_rate_vi))
            for _, r in sample_df.iterrows()
        ]
        sample_df["Food Grains for Class - V (Kg.)"] = [
            round_special_num(r["No of student availing MDM Class - V"] * grain_rate_v)
            for _, r in sample_df.iterrows()
        ]
        sample_df["Food grains for Class - VI to VIII (Kg.)"] = [
            round_special_num(r["No of student availing MDM Class - VI to VIII"] * grain_rate_vi)
            for _, r in sample_df.iterrows()
        ]
        st.session_state.df = sample_df
        # Pre-set reference balances
        st.session_state.ob_cooking_v = 7606.50
        st.session_state.ob_cooking_vi = 120570.50
        st.session_state.ob_grains_v = 288.0
        st.session_state.ob_grains_vi = 1960.0
        st.session_state.rec_cooking_v = 0.0
        st.session_state.rec_cooking_vi = 0.0
        st.session_state.rec_grains_v = 0.0
        st.session_state.rec_grains_vi = 0.0
        st.session_state.hon_received = 0.0
        st.session_state.hon_expenditure = 0.0
        st.session_state.cooks_df = pd.DataFrame(DEFAULT_COOKS)
        st.rerun()

with btn_c2:
    if st.button("🗑️ Clear Daily Data"):
        st.session_state.df = pd.DataFrame(columns=COLUMNS)
        st.rerun()

# ---------------------------------------------------------------------------
# Step 1: Optional CSV Data Loader
# ---------------------------------------------------------------------------
with st.expander("Step 1: Paste CSV Data (Optional)", expanded=False):
    data_input = st.text_area(
        "Paste your CSV Data here and click Load",
        height=120,
        placeholder="Date,No of student availing MDM Class - V,No of student availing MDM Class - VI to VIII,Items served..."
    )
    if st.button("Load Pasted CSV"):
        if data_input:
            try:
                from io import StringIO
                new_df = pd.read_csv(StringIO(data_input.strip()))
                if "Food Grains for Class - V" in new_df.columns and "Food Grains for Class - V (Kg.)" not in new_df.columns:
                    new_df.rename(columns={
                        "Food Grains for Class - V": "Food Grains for Class - V (Kg.)",
                        "Food grains for Class - VI to VIII": "Food grains for Class - VI to VIII (Kg.)"
                    }, inplace=True)
                if 'Items served' in new_df.columns:
                    new_df['Items served'] = new_df.apply(
                        lambda r: ensure_compulsory_menu(
                            r.get('Items served', ''),
                            has_students=(float(r.get("No of student availing MDM Class - V", 0) or 0) > 0 or
                                          float(r.get("No of student availing MDM Class - VI to VIII", 0) or 0) > 0)
                        ),
                        axis=1
                    )
                st.session_state.df = new_df
                st.rerun()
            except Exception as e:
                st.error(f"Error parsing CSV: {e}")

# ---------------------------------------------------------------------------
# Step 2: Daily Register Editor & Live Calculation
# ---------------------------------------------------------------------------
st.subheader("Step 2: Daily Serving Register")
st.caption("🍛 **ভাত, ডাল, আলু** is compulsory and added automatically! You don't need to type it — just type extra vegetables/items (e.g. `dim`, `potol`, `begun`, `shak`).")

calc_col1, calc_col2 = st.columns([3, 7])
with calc_col1:
    if st.button("🧮 Auto-Calculate Daily Costs & Grains", key="calc_btn"):
        for idx, row in st.session_state.df.iterrows():
            try:
                v_stu = float(row.get("No of student availing MDM Class - V", 0) or 0)
                vi_stu = float(row.get("No of student availing MDM Class - VI to VIII", 0) or 0)
                d_str = str(row.get("Date", ""))
                if "Total" not in d_str and "Days" not in d_str:
                    c_calc = (v_stu * cost_rate_v) + (vi_stu * cost_rate_vi)
                    gv_calc = v_stu * grain_rate_v
                    gvi_calc = vi_stu * grain_rate_vi
                    st.session_state.df.at[idx, "Cooking cost for Class - V to VIII"] = round_special_num(c_calc)
                    st.session_state.df.at[idx, "Food Grains for Class - V (Kg.)"] = round_special_num(gv_calc)
                    st.session_state.df.at[idx, "Food grains for Class - VI to VIII (Kg.)"] = round_special_num(gvi_calc)
            except Exception:
                pass
        st.rerun()

edited_df = st.data_editor(
    st.session_state.df,
    num_rows="dynamic",
    use_container_width=True,
    key="daily_editor"
)

# Intercept updates: preserve manual cell edits and auto-calculate rounded figures when students change
if not edited_df.equals(st.session_state.df):
    for idx in edited_df.index:
        old_v = st.session_state.df.at[idx, "No of student availing MDM Class - V"] if (idx in st.session_state.df.index and "No of student availing MDM Class - V" in st.session_state.df.columns) else None
        old_vi = st.session_state.df.at[idx, "No of student availing MDM Class - VI to VIII"] if (idx in st.session_state.df.index and "No of student availing MDM Class - VI to VIII" in st.session_state.df.columns) else None
        new_v = edited_df.at[idx, "No of student availing MDM Class - V"]
        new_vi = edited_df.at[idx, "No of student availing MDM Class - VI to VIII"]
        
        student_changed = (old_v != new_v) or (old_vi != new_vi)
        cost_val = edited_df.at[idx, "Cooking cost for Class - V to VIII"]
        
        try:
            v_val = float(new_v or 0)
        except (ValueError, TypeError):
            v_val = 0.0
        try:
            vi_val = float(new_vi or 0)
        except (ValueError, TypeError):
            vi_val = 0.0
        
        # Auto-compute rounded figures (no decimals except .5) in real-time
        if student_changed or pd.isna(cost_val) or cost_val == 0:
            if v_val > 0 or vi_val > 0:
                c_calc = (v_val * cost_rate_v) + (vi_val * cost_rate_vi)
                gv_calc = v_val * grain_rate_v
                gvi_calc = vi_val * grain_rate_vi
                edited_df.at[idx, "Cooking cost for Class - V to VIII"] = round_special_num(c_calc)
                edited_df.at[idx, "Food Grains for Class - V (Kg.)"] = round_special_num(gv_calc)
                edited_df.at[idx, "Food grains for Class - VI to VIII (Kg.)"] = round_special_num(gvi_calc)
            else:
                edited_df.at[idx, "Cooking cost for Class - V to VIII"] = 0
                edited_df.at[idx, "Food Grains for Class - V (Kg.)"] = 0
                edited_df.at[idx, "Food grains for Class - VI to VIII (Kg.)"] = 0
        # If student count didn't change but user modified cost/grains, preserve user's manual entry!
            
    if "Items served" in edited_df.columns:
        for idx in edited_df.index:
            try:
                v_s = float(edited_df.at[idx, "No of student availing MDM Class - V"] or 0)
            except (ValueError, TypeError):
                v_s = 0.0
            try:
                vi_s = float(edited_df.at[idx, "No of student availing MDM Class - VI to VIII"] or 0)
            except (ValueError, TypeError):
                vi_s = 0.0
            raw_item = edited_df.at[idx, "Items served"]
            edited_df.at[idx, "Items served"] = ensure_compulsory_menu(raw_item, has_students=(v_s > 0 or vi_s > 0))
    st.session_state.df = edited_df
    st.rerun()

# ---------------------------------------------------------------------------
# Live Metrics & Automatic Aggregations
# ---------------------------------------------------------------------------
# Filter clean daily rows (ignoring any user-typed summary row)
daily_mask = edited_df["Date"].notna() & (~edited_df["Date"].astype(str).str.contains("Total|Days|TOTAL", case=False))
clean_daily_df = edited_df[daily_mask].copy()

# Ensure numeric columns
for col in ["No of student availing MDM Class - V", "No of student availing MDM Class - VI to VIII", "Cooking cost for Class - V to VIII", "Food Grains for Class - V (Kg.)", "Food grains for Class - VI to VIII (Kg.)"]:
    if col in clean_daily_df.columns:
        clean_daily_df[col] = pd.to_numeric(clean_daily_df[col], errors="coerce").fillna(0)

active_days = len(clean_daily_df[clean_daily_df["Date"].astype(str).str.strip() != ""])
total_meals_v = int(clean_daily_df["No of student availing MDM Class - V"].sum())
total_meals_vi = int(clean_daily_df["No of student availing MDM Class - VI to VIII"].sum())

# Actual expenditure sums matching exact government MDM accounting rules (no decimals except .5):
calc_exp_cooking_v = round_special_num(total_meals_v * cost_rate_v)
calc_exp_cooking_vi = round_special_num(total_meals_vi * cost_rate_vi)
total_cooking_exp = round_special_num(calc_exp_cooking_v + calc_exp_cooking_vi)

calc_exp_grains_v = round_special_num(total_meals_v * grain_rate_v)
calc_exp_grains_vi = round_special_num(total_meals_vi * grain_rate_vi)
total_grains_exp = round_special_num(calc_exp_grains_v + calc_exp_grains_vi)

# Auto-sync Step 3 inputs whenever student meals change so expenditure inputs never stay stuck at 0.0:
if "prev_meals_v" not in st.session_state:
    st.session_state.prev_meals_v = total_meals_v
if "prev_meals_vi" not in st.session_state:
    st.session_state.prev_meals_vi = total_meals_vi

if total_meals_v != st.session_state.prev_meals_v or total_meals_vi != st.session_state.prev_meals_vi:
    st.session_state.exp_c_v_input = float(calc_exp_cooking_v)
    st.session_state.exp_c_vi_input = float(calc_exp_cooking_vi)
    st.session_state.exp_g_v_input = float(calc_exp_grains_v)
    st.session_state.exp_g_vi_input = float(calc_exp_grains_vi)
    st.session_state.prev_meals_v = total_meals_v
    st.session_state.prev_meals_vi = total_meals_vi

# Display Summary KPI Cards (No decimals except .5 in Streamlit!)
m_c1, m_c2, m_c3, m_c4, m_c5 = st.columns(5)
m_c1.metric("Serving Days", f"{active_days} days")
m_c2.metric("Class V Meals", f"{total_meals_v}")
m_c3.metric("Class VI-VIII Meals", f"{total_meals_vi}")
m_c4.metric("Cooking Cost Exp.", f"₹{format_val_with_unit(total_cooking_exp)}")
m_c5.metric("Grains Consumed", f"{format_val_with_unit(total_grains_exp)}")

# ---------------------------------------------------------------------------
# Step 3: Class-Wise Financial & Stock Accounting
# ---------------------------------------------------------------------------
st.markdown("---")
st.subheader("Step 3: Financial & Food Grains Accounting")
st.caption("Opening balances and funds received are tracked separately for Class V and Class VI–VIII, matching the official A4 Report and Utilization Certificate.")

# Default state initialization for financial balances
if "ob_cooking_v" not in st.session_state: st.session_state.ob_cooking_v = 7606.50
if "ob_cooking_vi" not in st.session_state: st.session_state.ob_cooking_vi = 120570.50
if "rec_cooking_v" not in st.session_state: st.session_state.rec_cooking_v = 0.0
if "rec_cooking_vi" not in st.session_state: st.session_state.rec_cooking_vi = 0.0

if "ob_grains_v" not in st.session_state: st.session_state.ob_grains_v = 288.0
if "ob_grains_vi" not in st.session_state: st.session_state.ob_grains_vi = 1960.0
if "rec_grains_v" not in st.session_state: st.session_state.rec_grains_v = 0.0
if "rec_grains_vi" not in st.session_state: st.session_state.rec_grains_vi = 0.0

if "hon_received" not in st.session_state: st.session_state.hon_received = 0.0
if "hon_expenditure" not in st.session_state: st.session_state.hon_expenditure = 0.0

fin_tab1, fin_tab2, fin_tab3 = st.tabs(["💵 Cooking Cost (Rs.)", "🌾 Food Grains - Rice (Kg.)", "👨‍🍳 Honorarium (Rs.)"])

with fin_tab1:
    f_c1, f_c2, f_c3 = st.columns(3)
    with f_c1:
        st.write("##### Class V (Primary)")
        ob_c_v = st.number_input("1. Opening Balance (Rs.)", value=float(st.session_state.ob_cooking_v), step=10.0, format="%.2f", key="ob_c_v_input")
        rec_c_v = st.number_input("3. Fund Received (Rs.)", value=float(st.session_state.rec_cooking_v), step=10.0, format="%.2f", key="rec_c_v_input")
        exp_c_v = st.number_input("8. Total Expenditure (Rs.)", value=float(calc_exp_cooking_v), step=10.0, format="%.2f", key="exp_c_v_input", help="Defaults to (Class V Meals * Rate)")
        cb_c_v = round(ob_c_v + rec_c_v - exp_c_v, 2)
        st.info(f"**Closing Balance V:** ₹{cb_c_v:,.2f}")
    with f_c2:
        st.write("##### Class VI to VIII (Upper Primary)")
        ob_c_vi = st.number_input("1. Opening Balance (Rs.)", value=float(st.session_state.ob_cooking_vi), step=10.0, format="%.2f", key="ob_c_vi_input")
        rec_c_vi = st.number_input("3. Fund Received (Rs.)", value=float(st.session_state.rec_cooking_vi), step=10.0, format="%.2f", key="rec_c_vi_input")
        exp_c_vi = st.number_input("8. Total Expenditure (Rs.)", value=float(round(calc_exp_cooking_vi)), step=10.0, format="%.2f", key="exp_c_vi_input", help="Defaults to (Class VI-VIII Meals * Rate)")
        cb_c_vi = round(ob_c_vi + rec_c_vi - exp_c_vi, 2)
        st.info(f"**Closing Balance VI–VIII:** ₹{cb_c_vi:,.2f}")
    with f_c3:
        st.write("##### Total Fund Summary")
        tot_ob_c = round(ob_c_v + ob_c_vi, 2)
        tot_rec_c = round(rec_c_v + rec_c_vi, 2)
        tot_exp_c = round(exp_c_v + exp_c_vi, 2)
        tot_cb_c = round(cb_c_v + cb_c_vi, 2)
        st.metric("Total Opening Fund", f"₹{tot_ob_c:,.2f}")
        st.metric("Total Received Fund", f"₹{tot_rec_c:,.2f}")
        st.metric("Total Cooking Exp.", f"₹{tot_exp_c:,.2f}")
        st.success(f"**Grand Closing Balance:** ₹{tot_cb_c:,.2f}")

with fin_tab2:
    g_c1, g_c2, g_c3 = st.columns(3)
    with g_c1:
        st.write("##### Class V (Primary)")
        ob_g_v = st.number_input("2. Opening Balance (Kg.)", value=float(st.session_state.ob_grains_v), step=1.0, format="%.2f", key="ob_g_v_input")
        rec_g_v = st.number_input("5. Grains Received (Kg.)", value=float(st.session_state.rec_grains_v), step=1.0, format="%.2f", key="rec_g_v_input")
        exp_g_v = st.number_input("9. Total Utilised (Kg.)", value=float(round(calc_exp_grains_v)), step=1.0, format="%.2f", key="exp_g_v_input", help="Defaults to (Class V Meals * 0.100 kg rounded)")
        cb_g_v = round(ob_g_v + rec_g_v - exp_g_v, 2)
        st.info(f"**Closing Grains V:** {cb_g_v:,.2f} Kg")
    with g_c2:
        st.write("##### Class VI to VIII (Upper Primary)")
        ob_g_vi = st.number_input("2. Opening Balance (Kg.)", value=float(st.session_state.ob_grains_vi), step=1.0, format="%.2f", key="ob_g_vi_input")
        rec_g_vi = st.number_input("5. Grains Received (Kg.)", value=float(st.session_state.rec_grains_vi), step=1.0, format="%.2f", key="rec_g_vi_input")
        exp_g_vi = st.number_input("9. Total Utilised (Kg.)", value=float(round(calc_exp_grains_vi)), step=1.0, format="%.2f", key="exp_g_vi_input", help="Defaults to (Class VI-VIII Meals * 0.150 kg rounded)")
        cb_g_vi = round(ob_g_vi + rec_g_vi - exp_g_vi, 2)
        st.info(f"**Closing Grains VI–VIII:** {cb_g_vi:,.2f} Kg")
    with g_c3:
        st.write("##### Total Food Grains Summary")
        tot_ob_g = round(ob_g_v + ob_g_vi, 2)
        tot_rec_g = round(rec_g_v + rec_g_vi, 2)
        tot_exp_g = round(exp_g_v + exp_g_vi, 2)
        tot_cb_g = round(cb_g_v + cb_g_vi, 2)
        st.metric("Total Opening Stock", f"{tot_ob_g:,.2f} Kg")
        st.metric("Total Grains Received", f"{tot_rec_g:,.2f} Kg")
        st.metric("Total Utilised Stock", f"{tot_exp_g:,.2f} Kg")
        st.success(f"**Grand Closing Stock:** {tot_cb_g:,.2f} Kg")

with fin_tab3:
    h_c1, h_c2, h_c3 = st.columns(3)
    with h_c1:
        hon_rec = st.number_input("4. Honorarium Received (Rs.)", value=float(st.session_state.hon_received), step=100.0, format="%.2f", key="hon_rec_input")
    with h_c2:
        hon_exp = st.number_input("10. Total Honorarium Expenditure (Rs.)", value=float(st.session_state.hon_expenditure), step=100.0, format="%.2f", key="hon_exp_input")
    with h_c3:
        hon_cb = round(hon_rec - hon_exp, 2)
        st.metric("Closing Honorarium Balance", f"₹{hon_cb:,.2f}")

    st.markdown("---")
    st.write("##### 👨‍🍳 Cook-cum-Helper Payment Roster (Table 4 in UC)")
    st.caption("Enter each cook's payment for the month. When funds are received/disbursed, set amounts below (e.g. ₹1,500 each). If no fund was received (like August 2026), keep 'NIL'.")

    p_col1, p_col2, p_col3, p_col4 = st.columns([2, 2, 2, 4])
    with p_col1:
        if st.button("⚡ Set ₹1,500 each", key="set_1500_btn"):
            st.session_state.cooks_df["Amount Received (In Rs.)"] = "1500"
            st.session_state.hon_expenditure = 1500.0 * len(st.session_state.cooks_df)
            st.rerun()
    with p_col2:
        if st.button("⚡ Set ₹2,000 each", key="set_2000_btn"):
            st.session_state.cooks_df["Amount Received (In Rs.)"] = "2000"
            st.session_state.hon_expenditure = 2000.0 * len(st.session_state.cooks_df)
            st.rerun()
    with p_col3:
        if st.button("⚡ Set NIL for all", key="set_nil_btn"):
            st.session_state.cooks_df["Amount Received (In Rs.)"] = "NIL"
            st.session_state.hon_expenditure = 0.0
            st.rerun()

    edited_cooks = st.data_editor(
        st.session_state.cooks_df,
        num_rows="dynamic",
        use_container_width=True,
        key="cooks_editor"
    )
    if not edited_cooks.equals(st.session_state.cooks_df):
        st.session_state.cooks_df = edited_cooks
        # If numeric amounts entered, auto-sync with hon_expenditure
        calc_tot_hon = 0.0
        for _, c_row in edited_cooks.iterrows():
            amt_val = str(c_row.get("Amount Received (In Rs.)", "")).replace("₹", "").replace("Rs", "").replace(",", "").strip()
            try:
                calc_tot_hon += float(amt_val)
            except ValueError:
                pass
        if calc_tot_hon > 0:
            st.session_state.hon_expenditure = calc_tot_hon
        st.rerun()

# ---------------------------------------------------------------------------
# Excel Generator (Strict A4 Scale & Structure matching Reference)
# ---------------------------------------------------------------------------
def generate_excel_report(
    clean_daily_df,
    month_year, school_name, enrol_v, enrol_vi,
    gram_sansad, circle_name, block_name, shg_bank_info,
    cost_rate_v, cost_rate_vi,
    ob_c_v, ob_c_vi, rec_c_v, rec_c_vi, exp_c_v, exp_c_vi,
    ob_g_v, ob_g_vi, rec_g_v, rec_g_vi, exp_g_v, exp_g_vi,
    hon_rec, hon_exp
):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "MDM Report"

    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.print_options.horizontalCentered = True
    ws.page_margins = PageMargins(left=0.25, right=0.25, top=0.5, bottom=0.5, header=0.3, footer=0.3)

    f_title = Font(name="Arial", size=9, bold=True)
    f_enroll = Font(name="Arial", size=8, bold=True)
    f_hdr9 = Font(name="Arial", size=9, bold=True)
    f_hdr8 = Font(name="Arial", size=8, bold=True)
    f_data_bold = Font(name="Calibri", size=11, bold=True)

    a_left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)
    a_center_wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)
    a_center = Alignment(horizontal="center", vertical="center")
    a_right = Alignment(horizontal="right", vertical="center")

    thin = Side(style="thin")
    full_border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for row in ws.iter_rows(min_row=1, max_row=48, min_col=1, max_col=8):
        for cell in row:
            cell.border = full_border

    for col in range(1, 9):
        ws.cell(row=49, column=col).border = Border(top=thin, left=thin if col==1 else None, right=thin if col==8 else None)
        ws.cell(row=50, column=col).border = Border(bottom=thin, left=thin if col==1 else None, right=thin if col==8 else None)

    ws.column_dimensions['A'].width = 5.0
    ws.column_dimensions['B'].width = 11.0
    ws.column_dimensions['C'].width = 11.45
    ws.column_dimensions['D'].width = 9.80
    ws.column_dimensions['E'].width = 28.90
    ws.column_dimensions['F'].width = 12.50
    ws.column_dimensions['G'].width = 14.50
    ws.column_dimensions['H'].width = 13.50

    row_heights = {
        1: 18.0, 2: 18.0, 3: 16.0, 4: 16.0,
        5: 14.5, 6: 14.5, 7: 14.5, 8: 14.5, 9: 14.5, 10: 14.5, 11: 14.5,
        12: 52.0,
        **{r: 14.55 for r in range(13, 39)},
        39: 16.0, 40: 20.0, 41: 26.0, 42: 16.0,
        43: 14.5, 44: 14.5, 45: 14.5, 46: 14.5, 47: 14.5, 48: 14.5,
        49: 71.4, 50: 30.0,
    }
    for r, h in row_heights.items():
        ws.row_dimensions[r].height = h

    merges = [
        "A1:H1", "A2:E2", "F2:G2", "A3:E3", "F3:G3",
        "B4:E4", "B5:E5", "B6:E6", "B7:E7", "B8:E8", "B9:E9", "B10:E10", "B11:E11",
        "A39:B39", "A40:H40", "A41:H41",
        "B42:E42", "B43:E43", "B44:E44", "B45:E45", "B46:E46", "B47:E47", "B48:E48",
        "A49:H49", "A50:D50", "E50:H50",
    ]
    for m in merges:
        ws.merge_cells(m)

    # Header Row 1-3
    c1 = ws.cell(row=1, column=1, value=f"MONTHLY REPORT FOR COOKING COST & FOOD GRAINS UNDER MID - DAY - MEAL FOR THE MONTH OF - {month_year}")
    c1.font = f_title; c1.alignment = a_left_wrap

    c2 = ws.cell(row=2, column=1, value=f"Name of the School : {school_name}")
    c2.font = f_title; c2.alignment = a_left_wrap
    c2_enrol = ws.cell(row=2, column=6, value=f"Total Enrolment : Class V = {enrol_v}")
    c2_enrol.font = f_enroll; c2_enrol.alignment = a_center_wrap

    c3 = ws.cell(row=3, column=1, value=f"Gram Sansad No. - {gram_sansad} , Circle : {circle_name}, Block : {block_name}")
    c3.font = f_title; c3.alignment = a_left_wrap
    c3_enrol = ws.cell(row=3, column=6, value=f"Class VI to VIII = {enrol_vi}")
    c3_enrol.font = f_hdr9; c3_enrol.alignment = a_center_wrap

    # Row 4: Top Particulars Header
    header_row4 = {1: "Sl.", 2: "Particular", 6: "Class -V", 7: "VI to VIII", 8: "Total"}
    for col, val in header_row4.items():
        cell = ws.cell(row=4, column=col, value=val)
        cell.font = f_hdr9; cell.alignment = a_center_wrap

    # Top Section Rows 5-11
    particulars_top = [
        ("Opening Balance Cooking Cost at the beginning of the month",
         to_excel_num(ob_c_v),
         to_excel_num(ob_c_vi),
         "=SUM(F5:G5)"),
        ("Opening Balance Food grains at the beginning of the month",
         to_excel_num(ob_g_v),
         to_excel_num(ob_g_vi),
         "=SUM(F6:G6)"),
        ("Fund for cooking cost received this month (for",
         to_excel_num(rec_c_v, is_nil_allowed=True),
         to_excel_num(rec_c_vi, is_nil_allowed=True),
         "NIL" if (rec_c_v == 0 and rec_c_vi == 0) else "=SUM(F7:G7)"),
        ("Honorarium received this month (for",
         "NIL",
         "NIL",
         to_excel_num(hon_rec, is_nil_allowed=True)),
        ("Food Grains received this month (for",
         to_excel_num(rec_g_v, is_nil_allowed=True),
         to_excel_num(rec_g_vi, is_nil_allowed=True),
         "NIL" if (rec_g_v == 0 and rec_g_vi == 0) else "=SUM(F9:G9)"),
        ("Total available Fund (1+3) (other grant)",
         to_excel_num(ob_c_v + rec_c_v),
         to_excel_num(ob_c_vi + rec_c_vi),
         "=SUM(F10:G10)"),
        ("Total Food Grains available (2+5) AMOUNT",
         to_excel_num(ob_g_v + rec_g_v),
         to_excel_num(ob_g_vi + rec_g_vi),
         "=SUM(F11:G11)")
    ]

    for idx, (title, val_v, val_vi, val_tot) in enumerate(particulars_top):
        r = 5 + idx
        ws.cell(row=r, column=1, value=idx + 1).font = f_data_bold
        ws.cell(row=r, column=1).alignment = a_center
        ws.cell(row=r, column=2, value=title).font = f_hdr9
        ws.cell(row=r, column=2).alignment = a_left_wrap
        
        c_v = ws.cell(row=r, column=6, value=val_v)
        c_vi = ws.cell(row=r, column=7, value=val_vi)
        c_tot = ws.cell(row=r, column=8, value=val_tot)
        for c in [c_v, c_vi, c_tot]:
            c.font = f_data_bold
            c.alignment = a_center

    # Row 12: Daily Register Headers
    headers = [
        "Sl. No.", "Date", "No of student\navailing MDM\nClass -V",
        "No of student\navailing MDM\nClass VI to VIII", "Items served",
        "Cooking cost\nfor Class - V to VIII", "Food Grains for\nClass - V (Kg.)",
        "Food grains for\nClass VI to VIII (Kg.)"
    ]
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=12, column=col_num, value=header)
        cell.font = f_hdr8; cell.alignment = a_center_wrap

    # Rows 13-38: Daily Data
    for i in range(26):
        r = 13 + i
        ws.cell(row=r, column=1, value=i + 1).font = f_hdr9
        ws.cell(row=r, column=1).alignment = a_center_wrap
        for col in range(2, 9):
            ws.cell(row=r, column=col).font = f_data_bold
            ws.cell(row=r, column=col).alignment = a_center_wrap

    last_row = 12
    row_count = 0
    for _, row in clean_daily_df.iterrows():
        current_row = 13 + row_count
        if current_row >= 39: break
        last_row = current_row
        row_count += 1

        v_stu = float(row.get("No of student availing MDM Class - V", 0) or 0)
        vi_stu = float(row.get("No of student availing MDM Class - VI to VIII", 0) or 0)

        # Ensure daily cost and grains take value from Streamlit table (or compute from rates if empty)
        raw_cost = float(row.get("Cooking cost for Class - V to VIII", 0) or 0)
        if raw_cost == 0 and (v_stu > 0 or vi_stu > 0):
            raw_cost = (v_stu * cost_rate_v) + (vi_stu * cost_rate_vi)
            
        raw_gv = float(row.get("Food Grains for Class - V (Kg.)", 0) or 0)
        if raw_gv == 0 and v_stu > 0:
            raw_gv = v_stu * grain_rate_v
            
        raw_gvi = float(row.get("Food grains for Class - VI to VIII (Kg.)", 0) or 0)
        if raw_gvi == 0 and vi_stu > 0:
            raw_gvi = vi_stu * grain_rate_vi

        ws.cell(row=current_row, column=1, value=row_count)
        ws.cell(row=current_row, column=2, value=str(row["Date"]))
        ws.cell(row=current_row, column=3, value=int(v_stu))
        ws.cell(row=current_row, column=4, value=int(vi_stu))
        ws.cell(row=current_row, column=5, value=ensure_compulsory_menu(row.get("Items served", ""), has_students=(v_stu > 0 or vi_stu > 0)))
        ws.cell(row=current_row, column=6, value=to_excel_num(raw_cost))
        ws.cell(row=current_row, column=7, value=to_excel_num(raw_gv))
        ws.cell(row=current_row, column=8, value=to_excel_num(raw_gvi))

    # Row 39: TOTAL Row
    ws.cell(row=39, column=1, value="TOTAL").font = f_hdr9
    ws.cell(row=39, column=1).alignment = a_center_wrap
    tot_meals_v = int(clean_daily_df["No of student availing MDM Class - V"].sum())
    tot_meals_vi = int(clean_daily_df["No of student availing MDM Class - VI to VIII"].sum())
    ws.cell(row=39, column=3, value=tot_meals_v).font = f_data_bold
    ws.cell(row=39, column=3).alignment = a_center
    ws.cell(row=39, column=4, value=tot_meals_vi).font = f_data_bold
    ws.cell(row=39, column=4).alignment = a_center
    
    # Calculate cooking cost sum matching monthly attendance & Step 3:
    if (exp_c_v + exp_c_vi) == 0:
        exp_c_v = round_special_num(tot_meals_v * cost_rate_v)
        exp_c_vi = round_special_num(tot_meals_vi * cost_rate_vi)
        tot_c_cost_num = exp_c_v + exp_c_vi
    else:
        tot_c_cost_num = exp_c_v + exp_c_vi
        
    if (exp_g_v + exp_g_vi) == 0:
        g_v_sum = clean_daily_df["Food Grains for Class - V (Kg.)"].sum() if "Food Grains for Class - V (Kg.)" in clean_daily_df.columns else (tot_meals_v * 0.100)
        g_vi_sum = clean_daily_df["Food grains for Class - VI to VIII (Kg.)"].sum() if "Food grains for Class - VI to VIII (Kg.)" in clean_daily_df.columns else (tot_meals_vi * 0.150)
        exp_g_v = round_special_num(g_v_sum)
        exp_g_vi = round_special_num(g_vi_sum)

    ws.cell(row=39, column=6, value=to_excel_num(tot_c_cost_num)).font = f_data_bold
    ws.cell(row=39, column=6).alignment = a_center
    
    # Grains totals (Formulas matching Completed_MDM_Report (1).xlsx)
    last_daily_row = 12 + row_count
    ws.cell(row=39, column=7, value=f"=SUM(G13:G{last_daily_row})").font = f_data_bold
    ws.cell(row=39, column=7).alignment = a_center
    ws.cell(row=39, column=8, value=f"=SUM(H13:H{last_daily_row})").font = f_data_bold
    ws.cell(row=39, column=8).alignment = a_center

    # Row 40: Formula Note
    formula_note = f"*(Total student of Class - V X {cost_rate_v:.1f}) + Total student of Class - VI to VIII X {cost_rate_vi:.2f})"
    c40 = ws.cell(row=40, column=1, value=formula_note)
    c40.font = f_hdr8; c40.alignment = a_left_wrap

    # Row 41: SHG & Bank Note
    c41 = ws.cell(row=41, column=1, value=shg_bank_info)
    c41.font = f_hdr8; c41.alignment = a_left_wrap

    # Row 42: Bottom Particulars Header
    header_row42 = {1: "Sl.", 2: "Particular", 6: "Class - V", 7: "VI to VIII", 8: "Total"}
    for col, val in header_row42.items():
        cell = ws.cell(row=42, column=col, value=val)
        cell.font = f_hdr9; cell.alignment = a_center_wrap

    # Bottom Section Rows 43-48
    particulars_bottom = [
        ("Total expenditure (Cooking Cost)",
         to_excel_num(exp_c_v),
         to_excel_num(exp_c_vi),
         "=SUM(F43:G43)"),
        ("Total Utilised (Food grains) in Kg.",
         to_excel_num(exp_g_v),
         to_excel_num(exp_g_vi),
         "=SUM(F44:G44)"),
        ("Total expenditure (Honorarium)",
         0,
         0,
         to_excel_num(hon_exp)),
        ("Closing Balance Fund (7-8)",
         to_excel_num(ob_c_v + rec_c_v - exp_c_v),
         to_excel_num(ob_c_vi + rec_c_vi - exp_c_vi),
         "=SUM(F46:G46)"),
        ("Closing Balance Food grains (7-9)",
         to_excel_num(ob_g_v + rec_g_v - exp_g_v),
         to_excel_num(ob_g_vi + rec_g_vi - exp_g_vi),
         "=SUM(F47:G47)"),
        ("Closing Balance Honorarium (4-10)",
         0,
         0,
         to_excel_num(max(0, hon_rec - hon_exp)))
    ]

    for idx, (title, val_v, val_vi, val_tot) in enumerate(particulars_bottom):
        r = 43 + idx
        ws.cell(row=r, column=1, value=idx + 8).font = f_hdr9
        ws.cell(row=r, column=1).alignment = a_center
        ws.cell(row=r, column=2, value=title).font = f_hdr9
        ws.cell(row=r, column=2).alignment = a_left_wrap
        
        c_v = ws.cell(row=r, column=6, value=val_v)
        c_vi = ws.cell(row=r, column=7, value=val_vi)
        c_tot = ws.cell(row=r, column=8, value=val_tot)
        for c in [c_v, c_vi, c_tot]:
            c.font = f_data_bold
            c.alignment = a_center

    # Row 50: Signatures
    c50_nodal = ws.cell(row=50, column=1, value="Signature of MDM Nodal Teacher")
    c50_nodal.font = f_hdr9; c50_nodal.alignment = a_center
    c50_hm = ws.cell(row=50, column=5, value="Signature of H M with seal")
    c50_hm.font = f_hdr9; c50_hm.alignment = a_right

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output.getvalue()

# ---------------------------------------------------------------------------
# Word MDCF / Utilization Certificate Generator
# ---------------------------------------------------------------------------
def set_docx_cell(cell, text, font_name="Arial", font_size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold

def generate_uc_docx(
    template_path,
    month_year, school_name, udise_code, state_name, district_name, block_name, village_name,
    enrol_v, enrol_vi, active_days, total_meals_v, total_meals_vi,
    ob_c_v, ob_c_vi, rec_c_v, rec_c_vi, exp_c_v, exp_c_vi,
    ob_g_v, ob_g_vi, rec_g_v, rec_g_vi, exp_g_v, exp_g_vi,
    hon_rec, hon_exp,
    cooks_df=None
):
    doc = docx.Document(template_path)

    # Table 0: School Details
    t0 = doc.tables[0]
    set_docx_cell(t0.rows[0].cells[1], month_year, bold=True)
    set_docx_cell(t0.rows[0].cells[4], udise_code, bold=True)
    set_docx_cell(t0.rows[0].cells[8], school_name, bold=True)
    
    set_docx_cell(t0.rows[3].cells[0], f"State / UT- {state_name}", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_docx_cell(t0.rows[3].cells[2], f"District- {district_name}", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_docx_cell(t0.rows[3].cells[7], block_name, bold=True)
    set_docx_cell(t0.rows[3].cells[9], f"Village/Ward- {village_name}", align=WD_ALIGN_PARAGRAPH.LEFT)
    set_docx_cell(t0.rows[4].cells[7], f"V - {enrol_v} / VI - VIII - {enrol_vi}", bold=True)

    # Table 1: Meals Availed Status
    t1 = doc.tables[1]
    days_str = str(active_days) if active_days > 0 else "NIL"
    set_docx_cell(t1.rows[1].cells[2], days_str)
    set_docx_cell(t1.rows[1].cells[3], days_str)
    set_docx_cell(t1.rows[2].cells[2], days_str)
    set_docx_cell(t1.rows[2].cells[3], days_str)
    set_docx_cell(t1.rows[3].cells[2], str(total_meals_v) if total_meals_v > 0 else "NIL", bold=True)
    set_docx_cell(t1.rows[3].cells[3], str(total_meals_vi) if total_meals_vi > 0 else "NIL", bold=True)

    # Table 2: Fund Details (Cooking Cost)
    t2 = doc.tables[2]
    # Primary (Row 2)
    cb_c_v = round_special_num(ob_c_v + rec_c_v - exp_c_v)
    set_docx_cell(t2.rows[2].cells[1], format_val_with_unit(ob_c_v, "Rs"))
    set_docx_cell(t2.rows[2].cells[2], format_val_with_unit(rec_c_v, "Rs", is_nil_allowed=True))
    set_docx_cell(t2.rows[2].cells[3], format_val_with_unit(exp_c_v, "Rs"))
    set_docx_cell(t2.rows[2].cells[4], format_val_with_unit(cb_c_v, "Rs"))

    # Upper Primary (Row 3)
    cb_c_vi = round_special_num(ob_c_vi + rec_c_vi - exp_c_vi)
    set_docx_cell(t2.rows[3].cells[1], format_val_with_unit(ob_c_vi, "Rs"))
    set_docx_cell(t2.rows[3].cells[2], format_val_with_unit(rec_c_vi, "Rs", is_nil_allowed=True))
    set_docx_cell(t2.rows[3].cells[3], format_val_with_unit(exp_c_vi, "Rs"))
    set_docx_cell(t2.rows[3].cells[4], format_val_with_unit(cb_c_vi, "Rs"))

    # Cook Cum Helper (Row 4)
    if hon_rec > 0 or hon_exp > 0:
        set_docx_cell(t2.rows[4].cells[2], format_val_with_unit(hon_rec, "Rs", is_nil_allowed=True))
        set_docx_cell(t2.rows[4].cells[3], format_val_with_unit(hon_exp, "Rs", is_nil_allowed=True))
        set_docx_cell(t2.rows[4].cells[4], format_val_with_unit(max(0.0, hon_rec - hon_exp), "Rs", is_nil_allowed=True))

    # Table 3: Cook Cum Helper Payment Details
    t3 = doc.tables[3]
    if cooks_df is not None and not cooks_df.empty:
        for i, (_, c_row) in enumerate(cooks_df.iterrows()):
            r_idx = i + 1
            if r_idx < len(t3.rows):
                set_docx_cell(t3.rows[r_idx].cells[1], str(c_row.get("Cook Name", "")), align=WD_ALIGN_PARAGRAPH.LEFT)
                set_docx_cell(t3.rows[r_idx].cells[2], str(c_row.get("Gender (M/F)", "")))
                set_docx_cell(t3.rows[r_idx].cells[3], str(c_row.get("Category (SC/ST/OBC/GEN)", "")))
                set_docx_cell(t3.rows[r_idx].cells[4], str(c_row.get("Payment Mode (Cash/Bank)", "")))
                raw_amt = str(c_row.get("Amount Received (In Rs.)", "NIL")).strip()
                amt_str = format_val_with_unit(raw_amt, "Rs", is_nil_allowed=True)
                set_docx_cell(t3.rows[r_idx].cells[5], amt_str, bold=True if amt_str != "NIL" else False)

    # Table 4: Food Grains Details (Rice)
    t4 = doc.tables[4]
    # Primary Rice (Row 4)
    cb_g_v = round_special_num(ob_g_v + rec_g_v - exp_g_v)
    set_docx_cell(t4.rows[4].cells[2], format_val_with_unit(ob_g_v, "Kg"))
    set_docx_cell(t4.rows[4].cells[3], format_val_with_unit(rec_g_v, "Kg", is_nil_allowed=True))
    set_docx_cell(t4.rows[4].cells[4], format_val_with_unit(exp_g_v, "Kg"))
    set_docx_cell(t4.rows[4].cells[5], format_val_with_unit(cb_g_v, "Kg"))

    # Upper Primary Rice (Row 6)
    cb_g_vi = round_special_num(ob_g_vi + rec_g_vi - exp_g_vi)
    set_docx_cell(t4.rows[6].cells[2], format_val_with_unit(ob_g_vi, "Kg"))
    set_docx_cell(t4.rows[6].cells[3], format_val_with_unit(rec_g_vi, "Kg", is_nil_allowed=True))
    set_docx_cell(t4.rows[6].cells[4], format_val_with_unit(exp_g_vi, "Kg"))
    set_docx_cell(t4.rows[6].cells[5], format_val_with_unit(cb_g_vi, "Kg"))

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

# ---------------------------------------------------------------------------
# Step 4: Export Reports & Utilization Certificates
# ---------------------------------------------------------------------------
st.markdown("---")
st.subheader("Step 4: Generate & Download Official Documents")
st.write("Generate both the **A4 Formatted MDM Monthly Report** and the **PM POSHAN Utilization Certificate / MDCF** in a single click.")

exp_col1, exp_col2, exp_col3 = st.columns(3)

# Build reports in-memory
clean_month_str = month_year.replace(" ", "_").replace("/", "-")
excel_bytes = generate_excel_report(
    clean_daily_df,
    month_year, school_name, enrol_v, enrol_vi,
    gram_sansad, circle_name, block_name, shg_bank_info,
    cost_rate_v, cost_rate_vi,
    ob_c_v, ob_c_vi, rec_c_v, rec_c_vi, exp_c_v, exp_c_vi,
    ob_g_v, ob_g_vi, rec_g_v, rec_g_vi, exp_g_v, exp_g_vi,
    hon_rec, hon_exp
)

template_doc_path = os.path.join(os.path.dirname(__file__), "mdm uc.docx")
docx_bytes = None
if os.path.exists(template_doc_path):
    docx_bytes = generate_uc_docx(
        template_doc_path,
        month_year, school_name, udise_code, state_name, district_name, block_name, village_name,
        enrol_v, enrol_vi, active_days, total_meals_v, total_meals_vi,
        ob_c_v, ob_c_vi, rec_c_v, rec_c_vi, exp_c_v, exp_c_vi,
        ob_g_v, ob_g_vi, rec_g_v, rec_g_vi, exp_g_v, exp_g_vi,
        hon_rec, hon_exp,
        cooks_df=st.session_state.cooks_df
    )

with exp_col1:
    st.download_button(
        label="📊 Download MDM Report (.xlsx)",
        data=excel_bytes,
        file_name=f"Completed_MDM_Report_{clean_month_str}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        type="primary",
        use_container_width=True
    )

with exp_col2:
    if docx_bytes:
        st.download_button(
            label="📝 Download PM POSHAN UC (.docx)",
            data=docx_bytes,
            file_name=f"PM_POSHAN_MDCF_UC_{clean_month_str}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
    else:
        st.warning("`mdm uc.docx` template not found in app directory.")

with exp_col3:
    if docx_bytes:
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"Completed_MDM_Report_{clean_month_str}.xlsx", excel_bytes)
            zf.writestr(f"PM_POSHAN_MDCF_UC_{clean_month_str}.docx", docx_bytes)
        zip_buf.seek(0)
        
        st.download_button(
            label="📦 Download Both as ZIP",
            data=zip_buf.getvalue(),
            file_name=f"MDM_Documents_{clean_month_str}.zip",
            mime="application/zip",
            use_container_width=True
        )
